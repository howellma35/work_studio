"""
多格式文件解析服务 — 支持 PDF、DOCX、TXT、MD、CSV、HTML
"""
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_file(file_path: Path, file_type: str) -> str:
    """
    解析文件并提取纯文本

    Args:
        file_path: 文件路径
        file_type: 文件扩展名（小写，不含点号）

    Returns:
        提取的文本内容
    """
    file_type = file_type.lower().strip(".")

    parsers = {
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "txt": _parse_text,
        "md": _parse_text,
        "markdown": _parse_text,
        "csv": _parse_csv,
        "html": _parse_html,
        "htm": _parse_html,
    }

    parser = parsers.get(file_type)
    if not parser:
        raise ValueError(f"不支持的文件格式: .{file_type}")

    text = parser(file_path)
    logger.info(f"Parsed {file_type} file: {file_path.name}, extracted {len(text)} chars")
    return text


def _parse_pdf(file_path: Path) -> str:
    """解析 PDF 文件"""
    import pdfplumber

    pages_text = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            # 提取表格
            tables = page.extract_tables()
            table_text = ""
            for table in tables:
                for row in table:
                    if row:
                        table_text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"

            pages_text.append(f"--- 第 {i + 1} 页 ---\n{page_text}\n{table_text}")

    return "\n\n".join(pages_text)


def _parse_docx(file_path: Path) -> str:
    """解析 DOCX 文件"""
    from docx import Document

    doc = Document(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)

    # 提取表格
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n\n".join(paragraphs)


def _parse_text(file_path: Path) -> str:
    """解析纯文本文件（TXT、MD）"""
    import chardet

    raw = file_path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        return raw.decode("utf-8", errors="replace")


def _parse_csv(file_path: Path) -> str:
    """解析 CSV 文件"""
    import csv
    import chardet

    raw = file_path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    text = raw.decode(encoding, errors="replace")
    lines = text.strip().split("\n")

    if not lines:
        return ""

    # 解析 CSV 并格式化为可读文本
    reader = csv.reader(lines)
    rows = list(reader)

    if len(rows) <= 1:
        return text

    # 第一行作为表头
    header = rows[0]
    result = [" | ".join(header), "-" * 40]

    for row in rows[1:]:
        result.append(" | ".join(row))

    return "\n".join(result)


def _parse_html(file_path: Path) -> str:
    """解析 HTML 文件"""
    from bs4 import BeautifulSoup
    import chardet

    raw = file_path.read_bytes()
    detected = chardet.detect(raw)
    encoding = detected.get("encoding", "utf-8") or "utf-8"

    html = raw.decode(encoding, errors="replace")
    soup = BeautifulSoup(html, "html.parser")

    # 移除 script 和 style 标签
    for tag in soup(["script", "style"]):
        tag.decompose()

    text = soup.get_text(separator="\n")
    # 清理多余空行
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)
